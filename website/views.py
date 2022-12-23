from django.shortcuts import render
import base64
import PyPDF2
import docx
import qrcode
from django.core.files.storage import default_storage
import os
import random
import string
from django.http import HttpResponse, JsonResponse
from docx2pdf import convert
from reportlab.pdfgen.canvas import Canvas


def generate_qr(content):
    data = str(content)
    qrcode_img = qrcode.make(data)
    img = base64.b64encode(qrcode_img.tobytes()).decode()
    # print(img.format())

    with default_storage.open(f'website/static/website/qrcodes/{data}.png', 'wb') as f:
        qrcode_img.save(f, 'PNG')

    return data


# class File_Convertion:

def pdf_to_word(file_path, ext):
    # Specify the input and output file paths
    input_file = file_path
    # output_file = f'website/static/website/converts/compressed.docx'

    doc = docx.Document()

    if ext == ".docx":

        with open('input.pdf', 'wb') as f:
            f.write(input_file.read())

        with open('input.pdf', 'rb') as f:
            pdf = PyPDF2.PdfFileReader(f)

            for page in range(pdf.getNumPages()):
                text = pdf.getPage(page).extractText()
                doc.add_paragraph(text)


        doc.save("output.docx")

        with open('output.docx', 'rb') as f:
            converted = f.read()
        #
        return converted

    if ext == ".pdf":

        document = docx.Document(file_path)
        paragraphs = [p.text for p in document.paragraphs]

        pdf_writer = PyPDF2.PdfFileWriter()
        page = pdf_writer.addBlankPage(width=595, height=842)
        canvas = Canvas(page)
        canvas.setFont("Helvetica", 12)
        for p in paragraphs:
            canvas.drawString(72, 756, p)
            canvas.drawString(72, 730, "-------------------")

        with open("document.pdf", "wb") as f:
            pdf_writer.write(f)

        with open("document.pdf", "rb") as f:
            to_pdf = f.read()

        return to_pdf


def generate_password(length):

    # Generate a random selection of letters, digits, and punctuation
    characters = string.ascii_letters + string.digits + string.punctuation

    # Shuffle the characters to ensure a random selection
    characters = list(characters)

    random.shuffle(characters)

    # Select the desired number of characters
    password = ''.join(characters[:length])

    return password


# ################################        VIEWS       ###############################################################
def index(request):
    context = {
    }
    return render(request, 'website/index.html', context)


def convert_file(request):
    if request.method == 'POST':

        try:
            file = request.FILES['file']
            file_ext = request.POST.get('extension')

            if file_ext == "pdf":
                ext = ".pdf"
            
            elif file_ext == "word":
                ext = ".docx"

            print(ext)

            word = pdf_to_word(file, ext)

            # Creating the HttpResponse object with the file content
            response = HttpResponse(word,
                                    content_type='application/docx')

            # Setting the Content-Disposition header to attachment
            response['Content-Disposition'] = 'attachment; filename=converted.docx'

            respond = {
                "response": response,
                "ext": ext,
            }

            return JsonResponse(respond)

        except FileNotFoundError:
            pass

    try:    
        os.remove('output.docx')
        os.remove('input.pdf')
        os.remove('output.pdf')
        os.remove('input.docx')
        os.remove('document.pdf')
        os.remove('document.docx')

    except FileNotFoundError:
        print("hey")

    return render(request, "website/convert.html")


def qrcode(request):
    qrcode_data = ''

    if request.method == "POST":
        qrcode_data = request.POST['qr']
        print(qrcode_data)

        generate_qr(qrcode_data)

    context = {
        'qrcode': qrcode_data,
    }

    return render(request, 'website/qrcode.html', context)


def p_generator(request):
    pwd = ''
    if request.method == 'POST':
        global length
        length = int(request.POST['length'])
        # global pwd
        pwd = generate_password(length)

    context = {
        'password': pwd
    }
    return render(request, "website/p_generator.html", context)


def about_us(request):
    return render(request, "website/about_us.html")
